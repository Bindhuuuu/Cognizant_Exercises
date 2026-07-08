"""
SkillGap AI - DOCX Resume Parser
Extracts raw text from DOCX resumes using python-docx.
"""

import io
from utils.logger import setup_logger

logger = setup_logger("docx_parser")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract all text from a DOCX file using python-docx.
    
    Args:
        file_content: Raw bytes of the DOCX file
    
    Returns:
        Extracted text as a single string
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"DOCX extracted {len(full_text)} characters from {len(doc.paragraphs)} paragraphs")
        return full_text
    
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""
